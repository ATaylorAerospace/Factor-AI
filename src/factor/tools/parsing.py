"""Document parsing tools for PDF and DOCX extraction."""

from __future__ import annotations

import logging
from pathlib import Path

from strands import tool

logger = logging.getLogger(__name__)


@tool
def parse_pdf(file_path: str) -> dict:
    """Extract text from a PDF document, preserving structure.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Dictionary with extracted text, page count, and metadata.
    """
    import fitz  # PyMuPDF

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}", "text": "", "pages": 0}

    doc = fitz.open(str(path))
    pages = []
    full_text_parts = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        pages.append({
            "page_number": page_num + 1,
            "text": text,
            "width": page.rect.width,
            "height": page.rect.height,
        })
        full_text_parts.append(text)

    doc.close()

    full_text = "\n\n".join(full_text_parts)
    logger.info("Parsed PDF %s: %d pages, %d characters", path.name, len(pages), len(full_text))

    return {
        "filename": path.name,
        "text": full_text,
        "pages": len(pages),
        "page_details": pages,
        "file_size_bytes": path.stat().st_size,
    }


@tool
def parse_docx(file_path: str) -> dict:
    """Extract text from a Word document, preserving structure.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Dictionary with extracted text, paragraph count, and metadata.
    """
    from docx import Document

    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}", "text": "", "paragraphs": 0}

    doc = Document(str(path))
    paragraphs = []
    full_text_parts = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "index": i,
                "text": text,
                "style": para.style.name if para.style else "Normal",
            })
            full_text_parts.append(text)

    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                tables_text.append(row_text)

    full_text = "\n\n".join(full_text_parts)
    if tables_text:
        full_text += "\n\n[TABLES]\n" + "\n".join(tables_text)

    logger.info("Parsed DOCX %s: %d paragraphs", path.name, len(paragraphs))

    return {
        "filename": path.name,
        "text": full_text,
        "paragraphs": len(paragraphs),
        "paragraph_details": paragraphs,
        "table_count": len(doc.tables),
        "file_size_bytes": path.stat().st_size,
    }
